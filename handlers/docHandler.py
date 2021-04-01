from docx import Document
from docx.enum.style import WD_STYLE

"""
    Processing Functions
"""

def get_document(path):
    return Document(path)

def get_body(document):
    return document.paragraphs

def get_table_contents(document):
    contents = []

    for table in document.tables:
        for i in range(0, len(table.rows)):
            for j in range(0, len(table.columns)):
                contents = contents + table.cell(i, j).paragraphs

    return contents

def get_header(document):
    return document.sections[0].header.paragraphs


#Returns the index of the string in the document
def get_text_position(string, paragraphs):
    idx = -1

    for i in range(0, len(paragraphs)):
        if paragraphs[i].text == string:
            idx = i

    return idx

#Return an array of paragraphs beginning at the first
#paragraph which matches the string, and ending at
#the next blank line
def get_section(string, paragraphs):
    section = []
    idx = get_text_position(string, paragraphs)

    while idx < len(paragraphs) and paragraphs[idx].text != "":
        section.append(paragraphs[idx])
        idx+=1

    return section

#Returns an array of elements which contain list items
#inside of an array of paragraph items. Only gets the
#first list.
def get_list(section):
    listStyles = [
        'List Item',
        'List Item (Inactive)',
        'Incomplete Item',
        'Complete Item'
    ]
    listStart = False
    list = []

    for graph in section:
        styleName = graph.style.name

        # See if the current graph is a list item
        if styleName in listStyles:
            listStart = True

            list.append(graph)
        else:
            # If the list has been started and we
            # hit a non-list item, leave the loop
            if listStart:
                break

    return list
